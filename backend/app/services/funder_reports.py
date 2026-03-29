import csv
import io
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta, timezone
from typing import Any

from docx import Document
from sqlalchemy import distinct, func, select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.appointment import Appointment
from app.models.client import Client
from app.models.follow_up import FollowUp
from app.models.service_entry import ServiceEntry


@dataclass
class FunderReportPayload:
    title: str
    org_name: str
    period_label: str
    start_date: date
    end_date: date
    metrics: dict[str, Any]
    raw_csv: str


def _day_bounds(value: date) -> tuple[datetime, datetime]:
    start = datetime.combine(value, time.min, tzinfo=timezone.utc)
    end = datetime.combine(value + timedelta(days=1), time.min, tzinfo=timezone.utc)
    return start, end


def _period_label(start_date: date, end_date: date) -> str:
    if start_date.year == end_date.year:
        return f"{start_date.strftime('%b %d')} to {end_date.strftime('%b %d, %Y')}"
    return f"{start_date.strftime('%b %d, %Y')} to {end_date.strftime('%b %d, %Y')}"


def _build_raw_csv(metrics: dict[str, Any]) -> str:
    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow(["section", "label", "value"])
    for key, value in metrics["summary"].items():
        writer.writerow(["summary", key, value])

    for item in metrics["service_breakdown"]:
        writer.writerow(["service_breakdown", item["service_type"], item["count"]])

    for item in metrics["languages_served"]:
        writer.writerow(["languages_served", item["language"], item["count"]])

    for item in metrics["monthly_services"]:
        writer.writerow(["monthly_services", item["month"], item["count"]])

    for item in metrics["appointments"]:
        writer.writerow(["appointments", item["status"], item["count"]])

    for item in metrics["follow_ups"]:
        writer.writerow(["follow_ups", item["status"], item["count"]])

    return output.getvalue()


async def build_funder_report_payload(
    db: AsyncSession,
    *,
    org_id: str,
    start_date: date,
    end_date: date,
) -> FunderReportPayload:
    if start_date > end_date:
        raise ValueError("Start date must be on or before end date.")

    start_dt, _ = _day_bounds(start_date)
    _, end_dt = _day_bounds(end_date)

    org_result = await db.execute(
        text("SELECT name FROM organizations WHERE id = :org_id"),
        {"org_id": org_id},
    )
    org_name = org_result.scalar_one_or_none() or "Organization"

    active_clients_result = await db.execute(
        select(func.count(Client.id)).where(
            Client.org_id == org_id,
            Client.status == "active",
        )
    )
    active_clients = int(active_clients_result.scalar() or 0)

    new_clients_result = await db.execute(
        select(func.count(Client.id)).where(
            Client.org_id == org_id,
            Client.created_at >= start_dt,
            Client.created_at < end_dt,
        )
    )
    new_clients = int(new_clients_result.scalar() or 0)

    services_result = await db.execute(
        select(
            func.count(ServiceEntry.id),
            func.count(distinct(ServiceEntry.client_id)),
        ).where(
            ServiceEntry.org_id == org_id,
            ServiceEntry.service_date >= start_date,
            ServiceEntry.service_date <= end_date,
        )
    )
    total_services, unique_clients_served = services_result.one()
    total_services = int(total_services or 0)
    unique_clients_served = int(unique_clients_served or 0)

    service_breakdown_result = await db.execute(
        select(ServiceEntry.service_type, func.count(ServiceEntry.id))
        .where(
            ServiceEntry.org_id == org_id,
            ServiceEntry.service_date >= start_date,
            ServiceEntry.service_date <= end_date,
        )
        .group_by(ServiceEntry.service_type)
        .order_by(func.count(ServiceEntry.id).desc(), ServiceEntry.service_type.asc())
    )
    service_breakdown = [
        {"service_type": service_type, "count": int(count)}
        for service_type, count in service_breakdown_result.all()
    ]

    languages_result = await db.execute(
        select(Client.language, func.count(distinct(Client.id)))
        .join(ServiceEntry, ServiceEntry.client_id == Client.id)
        .where(
            Client.org_id == org_id,
            ServiceEntry.org_id == org_id,
            ServiceEntry.service_date >= start_date,
            ServiceEntry.service_date <= end_date,
        )
        .group_by(Client.language)
        .order_by(func.count(distinct(Client.id)).desc(), Client.language.asc())
    )
    languages_served = [
        {"language": language or "unknown", "count": int(count)}
        for language, count in languages_result.all()
    ]

    monthly_services_result = await db.execute(
        select(
            func.date_trunc("month", ServiceEntry.service_date).label("service_month"),
            func.count(ServiceEntry.id),
        )
        .where(
            ServiceEntry.org_id == org_id,
            ServiceEntry.service_date >= start_date,
            ServiceEntry.service_date <= end_date,
        )
        .group_by("service_month")
        .order_by("service_month")
    )
    monthly_services = [
        {
            "month": service_month.date().isoformat() if hasattr(service_month, "date") else str(service_month),
            "count": int(count),
        }
        for service_month, count in monthly_services_result.all()
    ]

    appointments_result = await db.execute(
        select(Appointment.status, func.count(Appointment.id))
        .where(
            Appointment.org_id == org_id,
            Appointment.scheduled_at >= start_dt,
            Appointment.scheduled_at < end_dt,
        )
        .group_by(Appointment.status)
        .order_by(Appointment.status.asc())
    )
    appointments = [
        {"status": status, "count": int(count)}
        for status, count in appointments_result.all()
    ]

    pending_followups_result = await db.execute(
        select(func.count(FollowUp.id)).where(
            FollowUp.org_id == org_id,
            FollowUp.status == "pending",
        )
    )
    pending_follow_ups = int(pending_followups_result.scalar() or 0)

    followups_created_result = await db.execute(
        select(func.count(FollowUp.id)).where(
            FollowUp.org_id == org_id,
            FollowUp.created_at >= start_dt,
            FollowUp.created_at < end_dt,
        )
    )
    followups_created = int(followups_created_result.scalar() or 0)

    followups_completed_result = await db.execute(
        select(func.count(FollowUp.id)).where(
            FollowUp.org_id == org_id,
            FollowUp.completed_at.is_not(None),
            FollowUp.completed_at >= start_dt,
            FollowUp.completed_at < end_dt,
        )
    )
    followups_completed = int(followups_completed_result.scalar() or 0)

    metrics = {
        "summary": {
            "active_clients": active_clients,
            "new_clients": new_clients,
            "services_in_period": total_services,
            "unique_clients_served": unique_clients_served,
            "pending_follow_ups": pending_follow_ups,
            "follow_ups_created": followups_created,
            "follow_ups_completed": followups_completed,
        },
        "service_breakdown": service_breakdown,
        "languages_served": languages_served,
        "monthly_services": monthly_services,
        "appointments": appointments,
        "follow_ups": [
            {"status": "pending_now", "count": pending_follow_ups},
            {"status": "created_in_period", "count": followups_created},
            {"status": "completed_in_period", "count": followups_completed},
        ],
    }

    period_label = _period_label(start_date, end_date)
    title = f"{org_name} Funder Report ({period_label})"

    return FunderReportPayload(
        title=title,
        org_name=org_name,
        period_label=period_label,
        start_date=start_date,
        end_date=end_date,
        metrics=metrics,
        raw_csv=_build_raw_csv(metrics),
    )


def render_funder_report_docx(
    *,
    title: str,
    org_name: str,
    start_date: date,
    end_date: date,
    report_text: str,
    raw_csv: str | None = None,
) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"Organization: {org_name}")
    document.add_paragraph(
        f"Reporting period: {start_date.isoformat()} to {end_date.isoformat()}"
    )

    for paragraph in report_text.split("\n"):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())

    if raw_csv:
        document.add_heading("Appendix: Raw CSV Export", level=1)
        document.add_paragraph(raw_csv)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
